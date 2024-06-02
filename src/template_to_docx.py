from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from pymongo import MongoClient
from encrypt import decrypt_data


def tem1_to_docx(dict_):
    doc = Document()
    def add_paragraph(text, bold=False, align=None, font_size=14, space_after=Pt(0), space_before=Pt(0),
                      tab_stops=None):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(font_size)
        run.bold = bold
        if align:
            paragraph.alignment = align
        paragraph.space_after = space_after
        paragraph.space_before = space_before
        if tab_stops:
            for position, alignment in tab_stops.items():
                tab_stop = paragraph.paragraph_format.tab_stops.add_tab_stop(position, alignment)
        return paragraph
    ho_ten = dict_['Họ, tên người đề nghị thanh toán: '] if dict_['Họ, tên người đề nghị thanh toán: '] else ""
    don_vi = dict_['Đơn vị'] if dict_['Đơn vị'] else ""
    noi_dung_thanh_toan = dict_['Nội dung thanh toán'] if dict_['Nội dung thanh toán'] else ""
    ma_so_de_tai = dict_['Mã số để tài'] if dict_['Mã số để tài'] else ""
    so_tien = dict_['Số tiền'] if dict_['Số tiền'] else ""
    so_tien_bang_chu = dict_['Viết bằng chữ'] if dict_['Viết bằng chữ'] else ""
    header = doc.add_paragraph()
    header.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    header_run = header.add_run('BAN CƠ YẾU CHÍNH PHỦ                                                     \tMẫu số:')
    header_run.bold = True
    header_run.font.name = 'Times New Roman'
    header_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    header_run.font.size = Pt(14)

    header2 = doc.add_paragraph()
    header2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    header_run2 = header2.add_run('HỌC VIỆN KỸ THUẬT MẬT MÃ                                          \tSố: ')
    header_run2.bold = True
    header_run2.font.name = 'Times New Roman'
    header_run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    header_run2.font.size = Pt(14)

    # Thêm dòng kẻ
    add_paragraph('-----------------------------------------------------------------', align=WD_PARAGRAPH_ALIGNMENT.CENTER, font_size=14)

    # Thêm tiêu đề giấy đề nghị
    add_paragraph('GIẤY ĐỀ NGHỊ THANH TOÁN', bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER, font_size=14)
    add_paragraph('Ngày     tháng     năm 2023', align=WD_PARAGRAPH_ALIGNMENT.CENTER, font_size=14)

    # Thêm nội dung chính
    add_paragraph('Kính gửi:', bold=True, font_size=14)
    add_paragraph(f'Họ, tên người đề nghị thanh toán: {ho_ten}', font_size=14)
    add_paragraph(f'Đơn vị: {don_vi}', font_size=14)
    add_paragraph(f'Nội dung thanh toán: {noi_dung_thanh_toan}', font_size=14)
    add_paragraph(f'Mã số đề tài: {ma_so_de_tai} do: {ho_ten} làm chủ nhiệm.', font_size=14)
    add_paragraph(f'Số tiền: {so_tien}', font_size=14)
    add_paragraph(f'Viết bằng chữ: {so_tien_bang_chu}', font_size=14)
    add_paragraph('(Kèm theo … chứng từ gốc)', font_size=14)
    add_paragraph('Trong đó:', font_size=14)

    # Thêm danh sách các khoản chi
    # doc.add_paragraph('- Chi trực tiếp: 50.000.000 đồng').runs[0].font.size = Pt(14)
    # doc.add_paragraph('- Mua VTHH sử dụng ngay: 0 đồng').runs[0].font.size = Pt(14)
    # doc.add_paragraph('- Chi mua vật tư hàng hóa nhập kho: 0 đồng').runs[0].font.size = Pt(14)
    # doc.add_paragraph('- Chi mua tài sản cố định không qua lắp đặt, chạy thử: 0 đồng').runs[0].font.size = Pt(14)
    # doc.add_paragraph('- Chi mua tài sản cố định phải qua lắp đặt, chạy thử: 0 đồng').runs[0].font.size = Pt(14)
    # doc.add_paragraph('- Mua VTHH và TSCĐ cấp cho cấp dưới: 0 đồng').runs[0].font.size = Pt(14)

    # Thêm phần chữ ký
    add_paragraph('Người đề nghị thanh toán\t\tTrưởng phòng KHCN&HTPT\t\tTrưởng phòng Kế hoạch - Tài chính\t\tGiám đốc',
                tab_stops={Inches(1.5): WD_PARAGRAPH_ALIGNMENT.CENTER, Inches(3): WD_PARAGRAPH_ALIGNMENT.CENTER, Inches(4.5): WD_PARAGRAPH_ALIGNMENT.CENTER, Inches(6): WD_PARAGRAPH_ALIGNMENT.CENTER}, font_size=14)
    add_paragraph('(Ký, họ tên)\t\t\t\t\t(Ký, họ tên)\t\t\t\t\t(Ký, họ tên)\t\t\t\t\t(Ký, họ tên)',
                tab_stops={Inches(1.5): WD_PARAGRAPH_ALIGNMENT.CENTER, Inches(3): WD_PARAGRAPH_ALIGNMENT.CENTER, Inches(4.5): WD_PARAGRAPH_ALIGNMENT.CENTER, Inches(6): WD_PARAGRAPH_ALIGNMENT.CENTER}, font_size=14)

    # Lưu tài liệu
    doc.save('template.docx')


def tem3_to_docx(dict_):
    doc = Document()

    def add_paragraph(text, bold=False, align=None, font_size=14, space_after=Pt(0), space_before=Pt(0)):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(font_size)
        run.bold = bold
        if align:
            paragraph.alignment = align
        paragraph.space_after = space_after
        paragraph.space_before = space_before
        return paragraph

    def set_table_font(table, font_name='Times New Roman', font_size=10):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = font_name
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                        run.font.size = Pt(font_size)

    # Thông tin từ điển để lưu trữ các giá trị
    quyen_so = dict_['Quyển số'] if 'Quyển số' in dict_.keys() else ""
    no = dict_['Nợ'] if 'Nợ' in dict_.keys() else ""
    co = dict_['Có'] if 'Có' in dict_.keys() else ""
    hoten = dict_['Họ và tên người nhận tiền'] if 'Họ và tên người nhận tiền' in dict_.keys() else ""
    dia_chi = dict_['Địa chỉ'] if 'Địa chỉ' in dict_.keys() else ""
    ly_do = dict_['Lý do chi'] if 'Lý do chi' in dict_.keys() else ""
    so_tien = dict_['Số tiền'] if 'Số tiền' in dict_.keys() else ""
    bang_chu = dict_['Bằng chữ'] if 'Bằng chữ' in dict_.keys() else ""
    kem_theo = dict_['Kèm theo'] if 'Kèm theo' in dict_.keys() else ""
    danhan = dict_['Đã nhận đủ số tiền (viết bằng chữ)'] if 'Đã nhận đủ số tiền (viết bằng chữ)' in dict_.keys() else ""
    info = {
        "company_name": "CÔNG TY PHẦN MỀM QUẢN LÝ DOANH NGHIỆP (FAST)",
        "company_address": "Tầng 3, Tòa nhà CT1B - Khu VOV, Mễ Trì, Nam Từ Liêm, Hà Nội",
        "company_website": "www.fast.com.vn",
        "document_title": "PHIẾU CHI",
        "document_date": "Ngày 01 tháng 01 năm 2020",
        "receipt_book": "",
        # "receipt_number": "PC001",
        # "debit_account": "11211",
        # "credit_account": "1111",
        # "recipient_name": "Công ty CP SX - XD - TM ABC",
        # "recipient_address": "390 Điện Biên Phủ, Q.Bình Thạnh, TPHCM",
        # "reason_for_payment": "Nộp tiền mặt vào tài khoản ngân hàng.",
        # "amount": "120.000.000",
        # "amount_in_words": "Một trăm hai mươi triệu đồng",
        # "attached_documents": "0 Chứng từ gốc"
    }

    # Thêm thông tin tiêu đề
    add_paragraph(info["company_name"], font_size=14)
    add_paragraph(info["company_address"], font_size=14)
    add_paragraph(info["company_website"], font_size=14)
    add_paragraph('', font_size=14)  # Thêm dòng trống

    # Thêm tiêu đề Phiếu Chi
    add_paragraph(info["document_title"], bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER, font_size=14)
    add_paragraph(info["document_date"], align=WD_PARAGRAPH_ALIGNMENT.CENTER, font_size=14)

    # Thêm đoạn văn bản với tab stops căn phải
    right_aligned_paragraph = doc.add_paragraph()
    tab_stop = right_aligned_paragraph.paragraph_format.tab_stops.add_tab_stop(Pt(480), WD_TAB_ALIGNMENT.RIGHT,
                                                                               WD_TAB_LEADER.SPACES)

    # Tạo bảng thông tin số phiếu và tài khoản và chèn vào tài liệu
    table = doc.add_table(rows=4, cols=2)

    # Đặt văn bản trong bảng vào giữa
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            cell.paragraphs[0].space_after = Pt(0)  # Giảm khoảng cách sau đoạn văn
            cell.paragraphs[0].space_before = Pt(0)  # Giảm khoảng cách trước đoạn văn

    table.cell(0, 0).paragraphs[0].add_run('Quyển số: ').bold = True
    table.cell(0, 1).paragraphs[0].add_run(info["receipt_book"])

    table.cell(1, 0).paragraphs[0].add_run('Số: ').bold = True
    table.cell(1, 1).paragraphs[0].add_run(quyen_so)

    table.cell(2, 0).paragraphs[0].add_run('Nợ: ').bold = True
    table.cell(2, 1).paragraphs[0].add_run(no)

    table.cell(3, 0).paragraphs[0].add_run('Có: ').bold = True
    table.cell(3, 1).paragraphs[0].add_run(co)

    # Điều chỉnh phông chữ của bảng
    set_table_font(table, font_size=10)

    # Thêm đoạn văn sau bảng để đảm bảo vị trí
    doc.add_paragraph()

    # Thêm nội dung chi tiết
    add_paragraph(f'Họ và tên người nhận tiền: {hoten}', font_size=14)
    add_paragraph(f'Địa chỉ: {dia_chi}', font_size=14)
    add_paragraph(f'Lý do chi: {ly_do}', font_size=14)
    add_paragraph(f'Số tiền: {so_tien}', font_size=14)
    add_paragraph(f'Bằng chữ: {bang_chu}', font_size=14)
    add_paragraph(f'Kèm theo: {kem_theo}', font_size=14)

    # Thêm phần ký tên
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    table.cell(0, 0).paragraphs[0].add_run('GIÁM ĐỐC').bold = True
    table.cell(0, 1).paragraphs[0].add_run('KẾ TOÁN TRƯỞNG').bold = True
    table.cell(0, 2).paragraphs[0].add_run('THỦ QUỸ').bold = True
    table.cell(0, 3).paragraphs[0].add_run('NGƯỜI NHẬN TIỀN').bold = True

    # Thêm chữ ký
    signatures = [
        '(Ký, họ tên, đóng dấu)',
        '(Ký, họ tên)',
        '(Ký, họ tên)',
        '(Ký, họ tên)'
    ]
    for i, text in enumerate(signatures):
        paragraph = table.cell(1, i).paragraphs[0]
        run = paragraph.add_run(text)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(12)

    doc.add_paragraph('\n' * 3)  # Thêm 3 dòng trống

    # Thêm dòng xác nhận
    add_paragraph('Đã nhận đủ số tiền (viết bằng chữ): ', font_size=14)
    add_paragraph(danhan, font_size=14)

    # Lưu tài liệu
    doc.save('template.docx')