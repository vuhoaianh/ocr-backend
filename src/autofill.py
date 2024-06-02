import io
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from mongo_db_ultils import get_documents
import pymongo


def getIdentityInfo(user_id, document_id):
    cccd_info = get_documents('db_cccd', user_id, document_id)
    return cccd_info

def fill_docx_template(docx_path, variables):
    document = Document(docx_path)
    
    # Thiết lập font chữ và cỡ chữ cho toàn bộ văn bản
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
    
    # Điền các giá trị vào văn bản Word
    for paragraph in document.paragraphs:
        for key, value in variables.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in variables.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, value)
                            for run in paragraph.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(14)
    
    return document

def auto_fill_student_confirmation(template_id, id_card_data):
    if template_id == '1':
        docx_path = r"doc_temp\giay_xac_nhan_thanh_toan_template.docx"
    elif template_id == '4':
        docx_path = r"doc_temp\giay_xac_nhan_sinh_vien_template.docx"
    variables = {
        "name_position": id_card_data.get('Họ và tên', ""),
        "dob_position": id_card_data.get('Ngày sinh', ""),
        "gender_position": id_card_data.get('Giới tính', ""),
        "id_position": id_card_data.get('ID', ""),
        "permanent_address": id_card_data.get('Nơi thường trú', ""),
        # Các trường khác sẽ cần lấy từ nguồn dữ liệu sinh viên
        "phone_position": "",  # Cần điền thông tin phù hợp từ nguồn dữ liệu sinh viên
        "date_position": " ",  # Ngày cấp của CCCD/CMND
        "current_address": "",  # Nơi ở hiện nay của gia đình
        "type_pst": "",  # Loại hình đào tạo
        "edu_system": "",  # Hệ đào tạo
        "major_position": " ",  # Ngành học
        "class_position": "",  # Lớp học
        "id_number_position": "",  # Mã sinh viên
        "semester_position": "",  # Học kỳ
        "academic_year_position": "",  # Năm học
        "time_position": "",  # Thời gian khóa học
        "address_position": "",
        "topic_position": "",
        "money_position": "",
        "content_position": "",
        "money_text_position": "",
    }

    print(variables)

    filled_doc = fill_docx_template(docx_path, variables)
    filled_doc.save("template.docx")