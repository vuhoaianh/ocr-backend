import io

from docx import Document

def replace_variable_in_docx(docx_path, variables):
    document =  Document(io.BytesIO(docx_path))
    for paragraph in document.paragraphs:
        for key, value in variables.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in variables.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, value)

    return document

# Example usage:
docx_path = r"D:\IDE\backend-ocr-pdf\doc_temp\giay_xac_nhan_thanh_toan_config.docx"
variables = {
    "name_position": "John Doe",
    "Address_position": "123 Main Street",
    "content_position": "Payment for services rendered",
    "topic_position": "Project123",
    "money_position": "$1000",
    "money_text_position": "One thousand dollars"
} 

# Replace variables in the docx file
doc = replace_variable_in_docx(docx_path, variables)
# Save the modified docx file
doc.save("output_file.docx")
